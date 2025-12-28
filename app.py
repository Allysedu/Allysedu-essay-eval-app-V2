import streamlit as st
import pdfplumber
import io
import json
import pandas as pd
import re
from typing import List, Dict
from openai import OpenAI
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import zipfile
import os
from dotenv import load_dotenv
from difflib import SequenceMatcher
import matplotlib.pyplot as plt
import matplotlib
matplotlib.use('Agg')  # GUI 백엔드 사용 안 함
import numpy as np
import seaborn as sns

# .env 파일에서 환경 변수 로드
load_dotenv()

# ============================================
# 설정: 환경 변수 또는 .env 파일에서 로드
# ============================================
# API Key는 .env 파일에서 로드됩니다 (보안을 위해)
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")

# 관리자 계정 정보 (환경 변수에서 로드, 없으면 기본값 사용)
ADMIN_ID = os.getenv("ADMIN_ID", "ally365")
ADMIN_PASSWORD = os.getenv("ADMIN_PASSWORD", "angie1000")

# ============================================

# 페이지 설정
st.set_page_config(
    page_title="에세이 평가 앱",
    page_icon="📝",
    layout="wide"
)

# 세션 상태 초기화
# 기본 평가 기준 4개 설정
DEFAULT_CRITERIA = [
    {
        "name": "내용의 충실성",
        "description": "주제에 대한 이해도와 내용의 충실성을 평가합니다.",
        "max_score": 25.0,
        "min_score": 15.0,
        "weight": 1.0
    },
    {
        "name": "체계와 논리성",
        "description": "글의 구조와 논리적 전개를 평가합니다.",
        "max_score": 25.0,
        "min_score": 15.0,
        "weight": 1.0
    },
    {
        "name": "창의성과 노력",
        "description": "독창적인 관점과 노력의 흔적을 평가합니다.",
        "max_score": 25.0,
        "min_score": 15.0,
        "weight": 1.0
    },
    {
        "name": "윤리와 성실성",
        "description": "인용과 출처 표기의 정확성, 표절 여부를 평가합니다.",
        "max_score": 25.0,
        "min_score": 15.0,
        "weight": 1.0
    }
]

if 'evaluation_criteria' not in st.session_state:
    st.session_state.evaluation_criteria = DEFAULT_CRITERIA.copy()
if 'uploaded_pdfs' not in st.session_state:
    st.session_state.uploaded_pdfs = []
if 'extracted_texts' not in st.session_state:
    st.session_state.extracted_texts = []
if 'evaluation_results' not in st.session_state:
    st.session_state.evaluation_results = []
if 'is_logged_in' not in st.session_state:
    st.session_state.is_logged_in = False
if 'logged_in_user' not in st.session_state:
    st.session_state.logged_in_user = ""
if 'evaluation_year' not in st.session_state:
    st.session_state.evaluation_year = ""
if 'evaluation_semester' not in st.session_state:
    st.session_state.evaluation_semester = ""
if 'evaluation_subject' not in st.session_state:
    st.session_state.evaluation_subject = ""
if 'evaluation_title' not in st.session_state:
    st.session_state.evaluation_title = ""
if 'evaluated_essays' not in st.session_state:
    st.session_state.evaluated_essays = []
# 평가 기준 템플릿 파일 경로
CRITERIA_TEMPLATES_FILE = "saved_criteria_templates.json"

def load_criteria_templates() -> Dict:
    """저장된 평가 기준 템플릿을 파일에서 로드합니다."""
    if os.path.exists(CRITERIA_TEMPLATES_FILE):
        try:
            with open(CRITERIA_TEMPLATES_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception as e:
            st.error(f"평가 기준 템플릿 로드 중 오류 발생: {str(e)}")
            return {}
    return {}

def save_criteria_templates(templates: Dict):
    """평가 기준 템플릿을 파일에 저장합니다."""
    try:
        with open(CRITERIA_TEMPLATES_FILE, 'w', encoding='utf-8') as f:
            json.dump(templates, f, ensure_ascii=False, indent=2)
    except Exception as e:
        st.error(f"평가 기준 템플릿 저장 중 오류 발생: {str(e)}")

if 'saved_criteria_templates' not in st.session_state:
    # 파일에서 로드
    st.session_state.saved_criteria_templates = load_criteria_templates()
if 'selected_template' not in st.session_state:
    st.session_state.selected_template = None
if 'is_admin_logged_in' not in st.session_state:
    st.session_state.is_admin_logged_in = False
if 'allowed_users' not in st.session_state:
    st.session_state.allowed_users = {}  # {ID: {"name": 이름, "password": 비밀번호}}
if 'show_admin_mode' not in st.session_state:
    st.session_state.show_admin_mode = False
if 'adjusted_max_score' not in st.session_state:
    st.session_state.adjusted_max_score = None  # 사용자가 설정한 만점 (None이면 원래 점수 사용)
if 'show_accumulated' not in st.session_state:
    st.session_state.show_accumulated = False  # 누적 데이터 표시 여부

def extract_text_from_pdf(pdf_file) -> str:
    """PDF 파일에서 텍스트를 추출합니다."""
    try:
        text = ""
        with pdfplumber.open(io.BytesIO(pdf_file.read())) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n\n"
        return text
    except Exception as e:
        st.error(f"PDF 텍스트 추출 중 오류 발생: {str(e)}")
        return ""

def calculate_similarity(text1: str, text2: str) -> float:
    """두 텍스트 간의 유사도를 계산합니다 (0.0 ~ 1.0)."""
    # 공백과 줄바꿈 제거하여 비교
    text1_clean = re.sub(r'\s+', '', text1)
    text2_clean = re.sub(r'\s+', '', text2)
    
    if not text1_clean or not text2_clean:
        return 0.0
    
    # SequenceMatcher를 사용하여 유사도 계산
    similarity = SequenceMatcher(None, text1_clean, text2_clean).ratio()
    return similarity

def check_plagiarism(current_text: str, evaluated_essays: List[Dict]) -> Dict:
    """현재 에세이와 이전 평가된 에세이들의 유사도를 검사합니다."""
    if not evaluated_essays:
        return {
            "max_similarity": 0.0,
            "similar_essay": None,
            "plagiarism_detected": False,
            "similarity_percentage": 0.0
        }
    
    max_similarity = 0.0
    similar_essay = None
    
    for essay in evaluated_essays:
        similarity = calculate_similarity(current_text, essay.get('text', ''))
        if similarity > max_similarity:
            max_similarity = similarity
            similar_essay = essay.get('filename', '알 수 없음')
    
    similarity_percentage = max_similarity * 100
    
    return {
        "max_similarity": max_similarity,
        "similar_essay": similar_essay,
        "plagiarism_detected": similarity_percentage > 30.0,
        "similarity_percentage": similarity_percentage
    }

def evaluate_essay_with_ai(essay_text: str, criteria: List[Dict], api_key: str) -> Dict:
    """OpenAI API를 사용하여 에세이를 평가합니다."""
    try:
        client = OpenAI(api_key=api_key)
        
        # 평가 기준을 문자열로 변환
        criteria_text = ""
        for idx, criterion in enumerate(criteria, 1):
            description = criterion.get('description', '')
            criteria_text += f"{idx}. {criterion['name']}"
            if description:
                criteria_text += f" ({description})"
            criteria_text += f": 최저점 {criterion['min_score']}점, 최고점 {criterion['max_score']}점\n"
        
        # 프롬프트 작성
        system_prompt = """너는 전문 에세이 채점관이야. 사용자가 설정한 평가 기준과 배점을 바탕으로 업로드된 에세이를 분석해서 점수를 매기고 상세한 피드백을 제공해야 해.

평가할 때는:
1. 각 평가 기준 항목별로 정확하고 공정한 점수를 매겨야 해
2. 점수는 반드시 설정된 최저점과 최고점 범위 내에서 매겨야 해
3. 각 항목별로 왜 그 점수를 받았는지 매우 구체적이고 상세한 피드백을 한글로 제공해야 해
4. 에세이의 강점과 개선점을 명확히 지적해야 해
5. 각 항목별로 기술적 오류(맞춤법, 문법, 표현 오류 등)나 점수 하락 요인을 구체적으로 제시해야 해
6. 잘 작성한 부분은 반드시 강조하고 구체적인 예시를 들어 칭찬해야 해
7. 개선이 필요한 부분은 학생의 글에서 실제로 사용된 문장이나 표현을 예시로 들어 구체적으로 설명해야 해
8. 오류가 있는 경우 정확한 문장을 인용하고 올바른 표현을 제시해야 해
9. 전체적인 종합 평가도 포함해야 해

피드백은 다음 형식으로 작성해줘:
- 각 항목별 평가: "[항목명] (점수/최고점): [매우 구체적이고 상세한 평가 내용과 이유]. 

✨ 잘 작성한 점: [학생의 글에서 실제로 사용된 문장이나 표현을 예시로 들어 구체적으로 강조하고 칭찬]

⚠️ 개선할 점 및 오류: [학생의 글에서 실제로 사용된 문장이나 표현을 예시로 들어 구체적으로 지적. 오류가 있으면 정확한 문장을 인용하고 올바른 표현을 제시]"

- 종합 평가: "전체적으로 [종합적인 평가]"

중요: 
- 반드시 학생의 글에서 실제로 사용된 문장이나 표현을 예시로 들어야 해
- 잘 작성한 부분은 "✨ 잘 작성한 점:"으로 시작하고 구체적인 예시와 함께 강조해야 해
- 개선할 점은 "⚠️ 개선할 점 및 오류:"로 시작하고 실제 문장을 인용하여 구체적으로 설명해야 해
- 오류가 있으면 정확한 문장을 인용하고 올바른 표현을 제시해야 해

결과는 반드시 다음 JSON 형식으로 반환해야 해:
{
    "scores": {
        "항목명1": 점수(숫자),
        "항목명2": 점수(숫자),
        ...
    },
    "feedback": "상세한 피드백 내용 (각 항목별 평가와 종합 평가를 포함한 친절하고 구체적인 한글 피드백)"
}"""

        user_prompt = f"""다음은 평가 기준과 배점이야:

{criteria_text}

다음은 평가할 에세이 전문이야:

---
{essay_text}
---

위 에세이 전문을 읽고, 설정된 평가 기준과 배점에 따라 각 항목별로 점수를 매기고, 왜 그 점수를 받았는지 매우 구체적이고 상세한 피드백을 한글로 작성해줘. 

각 항목별로:
1. 잘 작성한 점은 학생의 글에서 실제로 사용된 문장이나 표현을 예시로 들어 구체적으로 강조하고 칭찬해줘
2. 개선할 점과 오류는 학생의 글에서 실제로 사용된 문장이나 표현을 예시로 들어 구체적으로 지적해줘
3. 오류가 있으면 정확한 문장을 인용하고 올바른 표현을 제시해줘
4. 전체적인 종합 평가도 포함해줘

반드시 학생의 글에서 실제로 사용된 문장이나 표현을 예시로 들어야 하며, 추상적인 설명보다는 구체적인 인용과 예시를 통해 설명해줘. JSON 형식으로 결과를 반환해줘."""

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            response_format={"type": "json_object"},
            temperature=0.3
        )
        
        # JSON 응답 파싱
        result = json.loads(response.choices[0].message.content)
        
        # 점수 검증 및 총점 계산 (가중치 반영)
        total_score = 0.0
        validated_scores = {}
        
        for criterion in criteria:
            criterion_name = criterion['name']
            score = result.get('scores', {}).get(criterion_name, 0.0)
            weight = criterion.get('weight', 1.0)  # 가중치 (기본값 1.0)
            
            # 점수가 범위 내에 있는지 확인
            if score < criterion['min_score']:
                score = criterion['min_score']
            elif score > criterion['max_score']:
                score = criterion['max_score']
            
            validated_scores[criterion_name] = float(score)
            # 가중치를 적용한 점수를 총점에 더함
            total_score += float(score) * float(weight)
        
        return {
            "scores": validated_scores,
            "total_score": total_score,
            "feedback": result.get('feedback', '피드백을 생성할 수 없습니다.')
        }
        
    except json.JSONDecodeError:
        st.error("AI 응답을 파싱하는 중 오류가 발생했습니다.")
        return None
    except Exception as e:
        error_str = str(e)
        # OpenAI API 429 에러 (Rate Limit 또는 잔액 부족) 처리
        if "429" in error_str or "rate limit" in error_str.lower() or "insufficient_quota" in error_str.lower():
            st.error("""
            ⚠️ **OpenAI API 사용량 초과 또는 잔액 부족**
            
            다음을 확인해주세요:
            1. OpenAI 계정에 충분한 크레딧이 있는지 확인
            2. API 사용량 한도를 초과하지 않았는지 확인
            3. 잠시 후 다시 시도해주세요
            
            OpenAI 대시보드에서 계정 상태를 확인하실 수 있습니다: https://platform.openai.com/usage
            """)
        else:
            st.error(f"AI 평가 중 오류 발생: {str(e)}")
        return None

def calculate_similarity(text1: str, text2: str) -> float:
    """두 텍스트 간의 유사도를 계산합니다 (0.0 ~ 1.0)."""
    # 공백과 줄바꿈 제거하여 비교
    text1_clean = re.sub(r'\s+', '', text1)
    text2_clean = re.sub(r'\s+', '', text2)
    
    if not text1_clean or not text2_clean:
        return 0.0
    
    # SequenceMatcher를 사용하여 유사도 계산
    similarity = SequenceMatcher(None, text1_clean, text2_clean).ratio()
    return similarity

def check_plagiarism(current_text: str, evaluated_essays: List[Dict]) -> Dict:
    """현재 에세이와 이전 평가된 에세이들의 유사도를 검사합니다."""
    if not evaluated_essays:
        return {
            "max_similarity": 0.0,
            "similar_essay": None,
            "plagiarism_detected": False,
            "similarity_percentage": 0.0
        }
    
    max_similarity = 0.0
    similar_essay = None
    
    for essay in evaluated_essays:
        similarity = calculate_similarity(current_text, essay.get('text', ''))
        if similarity > max_similarity:
            max_similarity = similarity
            similar_essay = essay.get('filename', '알 수 없음')
    
    similarity_percentage = max_similarity * 100
    
    return {
        "max_similarity": max_similarity,
        "similar_essay": similar_essay,
        "plagiarism_detected": similarity_percentage > 30.0,
        "similarity_percentage": similarity_percentage
    }

def evaluate_essay_with_plagiarism_check(essay_text: str, filename: str, criteria: List[Dict], api_key: str, evaluated_essays: List[Dict]) -> Dict:
    """표절 검사를 포함한 에세이 평가를 수행합니다."""
    # 표절 검사 수행
    plagiarism_result = check_plagiarism(essay_text, evaluated_essays)
    
    # AI 평가 수행
    evaluation_result = evaluate_essay_with_ai(essay_text, criteria, api_key)
    
    if not evaluation_result:
        return None
    
    # 평가기준 4번(윤리와 성실성)에 표절 검사 결과 반영
    ethics_criterion_name = "윤리와 성실성"
    
    # 평가 기준에서 "윤리와 성실성" 찾기
    ethics_criterion = None
    for criterion in criteria:
        if criterion['name'] == ethics_criterion_name:
            ethics_criterion = criterion
            break
    
    if ethics_criterion:
        similarity_percentage = plagiarism_result['similarity_percentage']
        original_score = evaluation_result['scores'].get(ethics_criterion_name, ethics_criterion['max_score'])
        
        # 표절 검사 결과에 따라 점수 조정
        if similarity_percentage >= 50.0:
            # 50% 이상 유사: 0점
            adjusted_score = 0.0
            plagiarism_message = f"⚠️ 표절 검사 결과: {similarity_percentage:.1f}% 유사도로 감지되어 0점 처리되었습니다."
            if plagiarism_result['similar_essay']:
                plagiarism_message += f" (유사 에세이: {plagiarism_result['similar_essay']})"
        elif similarity_percentage > 30.0:
            # 30% 초과: 10점
            adjusted_score = 10.0
            plagiarism_message = f"⚠️ 표절 검사 결과: {similarity_percentage:.1f}% 유사도로 감지되어 10점으로 조정되었습니다."
            if plagiarism_result['similar_essay']:
                plagiarism_message += f" (유사 에세이: {plagiarism_result['similar_essay']})"
        else:
            # 30% 이하: 원래 점수 유지
            adjusted_score = original_score
            plagiarism_message = f"✅ 표절 검사 결과: {similarity_percentage:.1f}% 유사도 (정상 범위)"
        
        # 점수 업데이트
        evaluation_result['scores'][ethics_criterion_name] = adjusted_score
        
        # 총점 재계산 (가중치 반영)
        total_score = 0.0
        for criterion in criteria:
            criterion_name = criterion['name']
            score = evaluation_result['scores'].get(criterion_name, 0.0)
            weight = criterion.get('weight', 1.0)
            total_score += float(score) * float(weight)
        
        evaluation_result['total_score'] = total_score
        
        # 피드백에 표절 검사 결과 추가
        if plagiarism_result['plagiarism_detected']:
            evaluation_result['feedback'] += f"\n\n【표절 검사 결과】\n{plagiarism_message}"
        else:
            evaluation_result['feedback'] += f"\n\n【표절 검사 결과】\n{plagiarism_message}"
        
        # 표절 검사 정보 저장
        evaluation_result['plagiarism_check'] = plagiarism_result
    
    return evaluation_result

def check_login(user_id: str, password: str) -> bool:
    """로그인 정보를 확인합니다."""
    # 관리자는 항상 로그인 가능
    if user_id == ADMIN_ID and password == ADMIN_PASSWORD:
        return True
    # 관리자가 추가한 사용자 확인
    user_info = st.session_state.allowed_users.get(user_id)
    if user_info and user_info.get("password") == password:
        return True
    return False

def check_admin_login(user_id: str, password: str) -> bool:
    """관리자 로그인 정보를 확인합니다."""
    return user_id == ADMIN_ID and password == ADMIN_PASSWORD

def parse_feedback(feedback_text: str, criteria: List[Dict]) -> Dict:
    """피드백 텍스트를 파싱하여 구조화된 데이터로 변환합니다."""
    feedback_lines = feedback_text.split('\n')
    feedback_data = {}
    general_feedback = []
    
    current_item = None
    current_content = []
    
    for line in feedback_lines:
        line = line.strip()
        if not line:
            continue
        
        # 항목명 패턴 찾기: "[항목명] (점수/최고점):" 형식
        item_pattern = r'\[([^\]]+)\]\s*\(([^)]+)\)\s*:\s*(.+)'
        match = re.match(item_pattern, line)
        
        if match:
            # 이전 항목 저장
            if current_item:
                feedback_data[current_item] = '\n'.join(current_content)
            
            # 새 항목 시작
            current_item = match.group(1)
            score_info = match.group(2)
            initial_content = match.group(3)
            current_content = [initial_content] if initial_content else []
        elif '종합' in line or '전체적으로' in line or '전체' in line:
            # 이전 항목 저장
            if current_item:
                feedback_data[current_item] = '\n'.join(current_content)
                current_item = None
                current_content = []
            general_feedback.append(line)
        elif current_item:
            # 현재 항목의 내용 추가
            current_content.append(line)
        else:
            # 항목명 패턴이 없으면 종합 평가로 처리
            if feedback_data:
                general_feedback.append(line)
    
    # 마지막 항목 저장
    if current_item:
        feedback_data[current_item] = '\n'.join(current_content)
    
    # 각 항목별로 잘 작성한 점과 개선할 점 추출
    structured_feedback = {}
    for criterion in criteria:
        criterion_name = criterion['name']
        item_feedback = feedback_data.get(criterion_name, "")
        
        # ✨ 잘 작성한 점 추출
        good_points = []
        # ⚠️ 개선할 점 및 오류 추출
        improvement_points = []
        # 일반 평가 내용
        general_item_feedback = []
        
        if item_feedback:
            lines = item_feedback.split('\n')
            current_section = None
            good_section_started = False
            improvement_section_started = False
            seen_good_points = set()  # 중복 방지용
            seen_improvement_points = set()  # 중복 방지용
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # ✨ 잘 작성한 점 섹션 시작 감지
                is_good_section_start = ('✨' in line or ('잘 작성한 점' in line and (':' in line or '：' in line)))
                # ⚠️ 개선할 점 및 오류 섹션 시작 감지
                is_improvement_section_start = ('⚠️' in line or (('개선할 점' in line or '오류' in line) and (':' in line or '：' in line)))
                
                # ✨ 잘 작성한 점 섹션 시작
                if is_good_section_start:
                    # 이미 good 섹션이 시작되었으면 이 줄은 무시 (중복 헤더 방지)
                    if good_section_started:
                        continue
                    # 새 섹션 시작
                    current_section = 'good'
                    good_section_started = True
                    improvement_section_started = False  # 다른 섹션 종료
                    # ✨ 또는 "잘 작성한 점:" 제거
                    clean_line = re.sub(r'^[✨\s]*잘\s*작성한\s*점\s*[:：]\s*', '', line, flags=re.IGNORECASE)
                    clean_line = clean_line.replace('✨', '').strip()
                    if clean_line and clean_line not in seen_good_points:
                        good_points.append(clean_line)
                        seen_good_points.add(clean_line)
                # ⚠️ 개선할 점 및 오류 섹션 시작
                elif is_improvement_section_start:
                    # 이미 improvement 섹션이 시작되었으면 이 줄은 무시 (중복 헤더 방지)
                    if improvement_section_started:
                        continue
                    # 새 섹션 시작
                    current_section = 'improvement'
                    improvement_section_started = True
                    good_section_started = False  # 다른 섹션 종료
                    # ⚠️ 또는 "개선할 점 및 오류:" 제거
                    clean_line = re.sub(r'^[⚠️\s]*개선할\s*점\s*(및\s*오류)?\s*[:：]\s*', '', line, flags=re.IGNORECASE)
                    clean_line = clean_line.replace('⚠️', '').strip()
                    if clean_line and clean_line not in seen_improvement_points:
                        improvement_points.append(clean_line)
                        seen_improvement_points.add(clean_line)
                # 현재 섹션에 내용 추가
                elif current_section == 'good' and good_section_started:
                    # 다른 섹션 시작 신호가 아니고, 중복이 아닌 경우만 추가
                    if line and not is_improvement_section_start and line not in seen_good_points:
                        good_points.append(line)
                        seen_good_points.add(line)
                elif current_section == 'improvement' and improvement_section_started:
                    # 다른 섹션 시작 신호가 아니고, 중복이 아닌 경우만 추가
                    if line and not is_good_section_start and line not in seen_improvement_points:
                        improvement_points.append(line)
                        seen_improvement_points.add(line)
                # 일반 평가 내용 (섹션 시작 전 또는 섹션 외)
                elif line and not is_good_section_start and not is_improvement_section_start:
                    # 중복 제거
                    if line not in general_item_feedback:
                        general_item_feedback.append(line)
        
        structured_feedback[criterion_name] = {
            'summary': '\n'.join(general_item_feedback) if general_item_feedback else item_feedback,
            'good_points': '\n'.join(good_points) if good_points else '',
            'improvement_points': '\n'.join(improvement_points) if improvement_points else ''
        }
    
    return {
        'items': structured_feedback,
        'general': '\n'.join(general_feedback) if general_feedback else ''
    }

def create_feedback_report(result: Dict, criteria: List[Dict], evaluation_info: Dict) -> BytesIO:
    """학생별 피드백 보고서를 Word 문서로 생성합니다."""
    doc = Document()
    
    # 페이지 방향을 가로(landscape)로 설정
    section = doc.sections[0]
    # A4 가로: 너비 11.69인치, 높이 8.27인치
    section.page_height = Inches(8.27)
    section.page_width = Inches(11.69)
    
    # 제목
    title = doc.add_heading('에세이 평가 보고서', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 평가 정보
    doc.add_heading('평가 정보', level=1)
    info_para = doc.add_paragraph()
    if evaluation_info.get('year'):
        info_para.add_run(f"평가 년도: {evaluation_info['year']}\n").bold = True
    if evaluation_info.get('semester'):
        info_para.add_run(f"학기: {evaluation_info['semester']}\n").bold = True
    if evaluation_info.get('subject'):
        info_para.add_run(f"과목명: {evaluation_info['subject']}\n").bold = True
    if evaluation_info.get('title'):
        info_para.add_run(f"평가 제목: {evaluation_info['title']}\n").bold = True
    
    # 학생 정보
    student_name = result['filename'].replace('.pdf', '').replace('.PDF', '')
    doc.add_heading('학생 정보', level=1)
    student_para = doc.add_paragraph()
    student_para.add_run(f"학생명: {student_name}\n").bold = True
    
    # 점수 요약 (가로형 표)
    doc.add_heading('점수 요약', level=1)
    
    # 점수 테이블 생성 (가로형: 헤더 행 + 데이터 행)
    table = doc.add_table(rows=2, cols=len(criteria) + 2)  # 평가 기준 열들 + 점수 열 + 총점 열
    table.style = 'Light Grid Accent 1'
    
    # 첫 번째 행: 헤더
    header_cells = table.rows[0].cells
    header_cells[0].text = '평가 기준'
    for idx, criterion in enumerate(criteria, 1):
        header_cells[idx].text = criterion['name']
    header_cells[-1].text = '총점'
    
    # 헤더 셀 굵게 표시
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # 두 번째 행: 점수 데이터
    data_cells = table.rows[1].cells
    data_cells[0].text = '점수'
    data_cells[0].paragraphs[0].runs[0].font.bold = True
    
    # 각 평가 기준별 점수 (가중치 반영)
    total_score = 0.0
    total_max = 0.0
    for idx, criterion in enumerate(criteria, 1):
        score = result['scores'].get(criterion['name'], 0.0)
        weight = criterion.get('weight', 1.0)
        data_cells[idx].text = f"{score:.1f} / {criterion['max_score']:.1f}"
        total_score += score * weight
        total_max += criterion['max_score'] * weight
    
    # 총점
    data_cells[-1].text = f"{total_score:.1f} / {total_max:.1f}"
    data_cells[-1].paragraphs[0].runs[0].font.bold = True
    
    # 상세 피드백
    doc.add_heading('상세 피드백', level=1)
    
    # 피드백 텍스트를 파싱하여 구조화된 데이터로 변환
    parsed_feedback = parse_feedback(result['feedback'], criteria)
    
    # 평가 기준별 피드백 표 생성
    # 종합의견 행을 포함하여 행 수 계산
    num_rows = len(criteria) + (1 if parsed_feedback['general'] else 0) + 1  # 헤더 + 기준 행 + 종합의견 행(있으면)
    feedback_table = doc.add_table(rows=num_rows, cols=2)
    feedback_table.style = 'Light Grid Accent 1'
    
    # 열 너비 설정 (평가 기준: 더 줄임, 상세 피드백: 넓게)
    feedback_table.columns[0].width = Inches(0.75)  # 평가 기준 열 (절반으로 줄임)
    feedback_table.columns[1].width = Inches(10.94)  # 상세 피드백 열 (넓게)
    
    # 헤더
    header_cells = feedback_table.rows[0].cells
    header_cells[0].text = '평가 기준'
    header_cells[1].text = '상세 피드백'
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # 각 평가 기준별 피드백 추가
    row_idx = 1
    for criterion in criteria:
        criterion_name = criterion['name']
        row_cells = feedback_table.rows[row_idx].cells
        row_cells[0].text = criterion_name
        
        # 해당 기준에 대한 구조화된 피드백 가져오기
        item_feedback = parsed_feedback['items'].get(criterion_name, {})
        summary = item_feedback.get('summary', '')
        good_points = item_feedback.get('good_points', '')
        improvement_points = item_feedback.get('improvement_points', '')
        
        # 피드백 내용 구성
        feedback_content = []
        if summary:
            feedback_content.append(f"【평가 요약】\n{summary}")
        if good_points:
            feedback_content.append(f"\n✨ 잘 작성한 점:\n{good_points}")
        if improvement_points:
            feedback_content.append(f"\n⚠️ 개선할 점 및 오류:\n{improvement_points}")
        
        row_cells[1].text = '\n'.join(feedback_content) if feedback_content else "피드백 없음"
        
        row_idx += 1
    
    # 종합의견 행 추가
    if parsed_feedback['general']:
        row_cells = feedback_table.rows[row_idx].cells
        row_cells[0].text = '종합의견'
        # 종합의견 셀을 굵게 표시
        for paragraph in row_cells[0].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        
        # 종합의견 내용
        row_cells[1].text = parsed_feedback['general']
        # 종합의견 내용도 굵게 표시
        for paragraph in row_cells[1].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # 문서를 BytesIO로 저장
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

def admin_mode():
    """관리자 모드 페이지"""
    st.title("👑 관리자 모드")
    st.markdown("---")
    
    if not st.session_state.is_admin_logged_in:
        st.markdown("### 관리자 로그인")
        
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            with st.form("admin_login_form"):
                admin_id = st.text_input("관리자 ID", placeholder="관리자 ID를 입력하세요", key="admin_id_input")
                admin_password = st.text_input("관리자 비밀번호", type="password", placeholder="관리자 비밀번호를 입력하세요", key="admin_password_input")
                
                admin_login_button = st.form_submit_button("관리자 로그인", type="primary", use_container_width=True)
                
                if admin_login_button:
                    if check_admin_login(admin_id, admin_password):
                        st.session_state.is_admin_logged_in = True
                        st.success("✅ 관리자 로그인 성공!")
                        st.rerun()
                    else:
                        st.error("❌ 관리자 ID 또는 비밀번호가 올바르지 않습니다.")
        
        st.markdown("---")
        if st.button("← 메인으로 돌아가기", use_container_width=True):
            st.session_state.show_admin_mode = False
            st.session_state.is_admin_logged_in = False
            st.rerun()
        return
    
    # 관리자 로그인된 상태
    st.success(f"✅ 관리자로 로그인됨: {ADMIN_ID}")
    
    if st.button("🚪 관리자 로그아웃", use_container_width=True):
        st.session_state.is_admin_logged_in = False
        st.rerun()
    
    st.markdown("---")
    st.header("👥 사용자 관리")
    
    # 현재 허용된 사용자 목록
    st.subheader("현재 허용된 사용자 목록")
    
    if st.session_state.allowed_users:
        user_data = {
            "이름": [user_info["name"] for user_info in st.session_state.allowed_users.values()],
            "ID": list(st.session_state.allowed_users.keys())
        }
        user_df = pd.DataFrame(user_data)
        st.dataframe(user_df, use_container_width=True, hide_index=True)
        
        # 사용자 삭제
        st.markdown("### 사용자 삭제")
        delete_user_id = st.selectbox(
            "삭제할 사용자 ID 선택",
            options=[""] + list(st.session_state.allowed_users.keys()),
            key="delete_user_select"
        )
        
        if delete_user_id and st.button("🗑️ 사용자 삭제", type="primary", use_container_width=True):
            deleted_name = st.session_state.allowed_users[delete_user_id]["name"]
            del st.session_state.allowed_users[delete_user_id]
            st.success(f"✅ '{deleted_name}' ({delete_user_id}) 사용자가 삭제되었습니다!")
            st.rerun()
    else:
        st.info("💡 등록된 사용자가 없습니다. 아래에서 사용자를 추가하세요.")
    
    st.markdown("---")
    
    # 새 사용자 추가
    st.subheader("새 사용자 추가")
    
    with st.form("add_user_form"):
        new_user_name = st.text_input("사용자 이름", placeholder="예: 홍길동", key="new_user_name")
        new_user_id = st.text_input("사용자 ID", placeholder="예: hong123", key="new_user_id")
        new_user_password = st.text_input("비밀번호", type="password", placeholder="비밀번호를 입력하세요", key="new_user_password")
        
        add_user_button = st.form_submit_button("➕ 사용자 추가", type="primary", use_container_width=True)
        
        if add_user_button:
            if not new_user_name or not new_user_id or not new_user_password:
                st.error("❌ 모든 항목을 입력해주세요.")
            elif new_user_id == ADMIN_ID:
                st.error("❌ 관리자 ID는 사용할 수 없습니다.")
            elif new_user_id in st.session_state.allowed_users:
                st.error("❌ 이미 존재하는 ID입니다.")
            else:
                st.session_state.allowed_users[new_user_id] = {
                    "name": new_user_name,
                    "password": new_user_password
                }
                st.success(f"✅ '{new_user_name}' ({new_user_id}) 사용자가 추가되었습니다!")
                st.rerun()
    
    st.markdown("---")
    if st.button("← 메인으로 돌아가기", use_container_width=True):
        st.session_state.show_admin_mode = False
        st.rerun()

def main():
    # 관리자 모드 체크
    if st.session_state.get('show_admin_mode', False):
        admin_mode()
        return
    
    # 로그인 체크
    if not st.session_state.is_logged_in:
        # 로그인 화면
        # 우측 상단에 관리자 모드 버튼
        col_left, col_right = st.columns([10, 1])
        with col_right:
            if st.button("👑 관리자", use_container_width=True, type="secondary"):
                st.session_state.show_admin_mode = True
                st.rerun()
        
        st.title("🔐 에세이 평가 앱 로그인")
        st.markdown("---")
        
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            st.markdown("### 로그인이 필요합니다")
            
            with st.form("login_form"):
                user_id = st.text_input("사용자 ID", placeholder="ID를 입력하세요")
                password = st.text_input("비밀번호", type="password", placeholder="비밀번호를 입력하세요")
                
                login_button = st.form_submit_button("로그인", type="primary", use_container_width=True)
                
                if login_button:
                    if check_login(user_id, password):
                        st.session_state.is_logged_in = True
                        st.session_state.logged_in_user = user_id
                        st.success("✅ 로그인 성공!")
                        st.rerun()
                    else:
                        st.error("❌ ID 또는 비밀번호가 올바르지 않습니다.")
        
        st.markdown("---")
        st.info("💡 허용된 사용자만 이 앱을 사용할 수 있습니다.")
        return
    
    # 로그인된 사용자만 접근 가능
    # 사이드바: 로그인 정보 및 사용 방법
    with st.sidebar:
        st.header("⚙️ 설정")
        st.markdown("---")
        
        st.success(f"✅ 로그인됨: {st.session_state.logged_in_user}")
        
        if st.button("🚪 로그아웃", use_container_width=True):
            st.session_state.is_logged_in = False
            st.session_state.logged_in_user = ""
            st.session_state.evaluation_criteria = DEFAULT_CRITERIA.copy()
            st.session_state.uploaded_pdfs = []
            st.session_state.extracted_texts = []
            st.session_state.evaluation_results = []
            st.rerun()
        
        st.markdown("---")
        st.markdown("### 📚 저장된 평가 기준")
        
        # 저장된 평가 기준 목록 표시
        if st.session_state.saved_criteria_templates:
            template_names = list(st.session_state.saved_criteria_templates.keys())
            # format_func을 사용하여 전체 텍스트가 보이도록 설정
            def format_template_name(option):
                if option == "":
                    return "선택하세요..."
                return option  # 전체 텍스트 반환
            
            selected_template_name = st.selectbox(
                "평가 기준 선택",
                options=[""] + template_names,
                key="template_selector",
                help="저장된 평가 기준을 선택하면 현재 평가 기준으로 복사됩니다.",
                format_func=format_template_name
            )
            
            # 선택된 항목의 전체 텍스트 표시
            if selected_template_name:
                st.caption(f"📌 선택된 평가 기준: **{selected_template_name}**")
            
            if selected_template_name and selected_template_name != st.session_state.selected_template:
                # 선택한 템플릿을 현재 평가 기준으로 복사
                import copy
                st.session_state.evaluation_criteria = copy.deepcopy(st.session_state.saved_criteria_templates[selected_template_name])
                st.session_state.selected_template = selected_template_name
                # 평가 제목도 업데이트
                st.session_state.evaluation_title = selected_template_name
                st.success(f"✅ '{selected_template_name}' 평가 기준이 적용되었습니다!")
                st.rerun()
        else:
            st.info("💡 저장된 평가 기준이 없습니다. 평가 기준을 설정하고 저장해보세요!")
        
        st.markdown("---")
        
        # 삭제 기능 (별도 섹션으로 분리)
        st.markdown("### 🗑️ 평가 기준 삭제")
        
        if st.session_state.saved_criteria_templates:
            # 삭제 모드 토글 버튼
            if 'delete_mode' not in st.session_state:
                st.session_state.delete_mode = False
            
            if st.button("🗑️ 삭제 모드", use_container_width=True, type="secondary"):
                st.session_state.delete_mode = not st.session_state.delete_mode
                st.rerun()
            
            if st.session_state.delete_mode:
                st.warning("⚠️ 삭제 모드가 활성화되었습니다.")
                
                template_names = list(st.session_state.saved_criteria_templates.keys())
                delete_template_name = st.selectbox(
                    "삭제할 평가 기준 선택",
                    options=[""] + template_names,
                    key="delete_template_selector",
                    help="삭제할 평가 기준을 선택하세요.",
                    format_func=format_template_name
                )
                
                if delete_template_name:
                    col1, col2 = st.columns(2)
                    with col1:
                        if st.button(f"✅ 삭제 확인", use_container_width=True, type="primary", key="confirm_delete"):
                            del st.session_state.saved_criteria_templates[delete_template_name]
                            # 현재 선택된 템플릿이 삭제된 경우 선택 해제
                            if st.session_state.selected_template == delete_template_name:
                                st.session_state.selected_template = None
                            # 파일에 저장 (삭제 반영)
                            save_criteria_templates(st.session_state.saved_criteria_templates)
                            st.session_state.delete_mode = False
                            st.success(f"✅ '{delete_template_name}' 평가 기준이 삭제되었습니다!")
                            st.rerun()
                    with col2:
                        if st.button("❌ 취소", use_container_width=True, key="cancel_delete"):
                            st.session_state.delete_mode = False
                            st.rerun()
        else:
            st.info("💡 삭제할 평가 기준이 없습니다.")
        
        st.markdown("---")
        st.markdown("### 📖 사용 방법")
        st.markdown("""
        1. 평가 제목 입력
        2. 평가 기준 설정
        3. PDF 파일 업로드
        4. 텍스트 추출
        5. 평가 실행
        """)
    
    # 우측 상단에 관리자 모드 버튼
    col_left, col_right = st.columns([10, 1])
    with col_right:
        if st.button("👑 관리자", use_container_width=True, type="secondary"):
            st.session_state.show_admin_mode = True
            st.rerun()
    
    st.title("📝 에세이 평가 앱")
    st.markdown("---")
    
    # 1. 평가 정보 입력
    st.header("1️⃣ 평가 정보")
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        # 위젯 key와 session_state를 분리하여 충돌 방지
        evaluation_year = st.text_input(
            "평가 년도",
            placeholder="예: 2024",
            key="widget_evaluation_year",
            value=st.session_state.evaluation_year if st.session_state.evaluation_year else ""
        )
        # 위젯 값이 변경되었을 때만 session_state 업데이트
        if evaluation_year != st.session_state.evaluation_year:
            st.session_state.evaluation_year = evaluation_year
    
    with col2:
        semester_options = ["", "1학기", "2학기", "여름학기", "겨울학기"]
        current_index = 0
        if st.session_state.evaluation_semester in semester_options:
            current_index = semester_options.index(st.session_state.evaluation_semester)
        
        evaluation_semester = st.selectbox(
            "학기",
            options=semester_options,
            key="widget_evaluation_semester",
            index=current_index
        )
        # 위젯 값이 변경되었을 때만 session_state 업데이트
        if evaluation_semester != st.session_state.evaluation_semester:
            st.session_state.evaluation_semester = evaluation_semester
    
    with col3:
        evaluation_subject = st.text_input(
            "과목명",
            placeholder="예: 영어작문, 국어 등",
            key="widget_evaluation_subject",
            value=st.session_state.evaluation_subject if st.session_state.evaluation_subject else ""
        )
        # 위젯 값이 변경되었을 때만 session_state 업데이트
        if evaluation_subject != st.session_state.evaluation_subject:
            st.session_state.evaluation_subject = evaluation_subject
    
    st.markdown("---")
    
    # 평가 제목 입력
    st.subheader("평가 제목")
    
    # 평가 제목 자동 생성 예시 표시
    auto_title = ""
    if evaluation_year and evaluation_semester and evaluation_subject:
        auto_title = f"{evaluation_year}년 {evaluation_semester} {evaluation_subject} 에세이 평가"
        st.info(f"💡 자동 생성 예시: **{auto_title}** (아래에서 직접 입력하세요)")
    
    # 사용자가 직접 입력
    col1, col2 = st.columns([4, 1])
    with col1:
        evaluation_title = st.text_input(
            "평가 제목을 입력하세요",
            placeholder="예: 2024년 1학기 영어작문 중간고사 에세이 평가",
            key="widget_evaluation_title",
            value=st.session_state.evaluation_title if st.session_state.evaluation_title else "",
            label_visibility="visible"
        )
    
    with col2:
        st.markdown("<br>", unsafe_allow_html=True)  # 버튼을 텍스트 입력과 같은 높이로 맞추기
        if st.button("✅ 확인", key="confirm_title", use_container_width=True, type="primary"):
            if evaluation_title:
                st.session_state.evaluation_title = evaluation_title
                st.success(f"✅ 평가 제목이 설정되었습니다: **{evaluation_title}**")
                st.rerun()
            else:
                st.warning("⚠️ 평가 제목을 입력해주세요.")
    
    # 현재 설정된 평가 제목 표시
    if st.session_state.evaluation_title:
        st.markdown(f"**현재 설정된 평가 제목:** {st.session_state.evaluation_title}")
    
    st.markdown("---")
    
    # 2. 평가 기준 설정
    st.header("2️⃣ 평가 기준 설정")
    
    # 기본값 설정: 저장된 기준이 없으면 기본 4개 사용
    if not st.session_state.evaluation_criteria or len(st.session_state.evaluation_criteria) == 0:
        st.session_state.evaluation_criteria = DEFAULT_CRITERIA.copy()
    
    num_criteria = st.number_input(
        "평가 기준 항목 개수 (최대 10개)",
        min_value=1,
        max_value=10,
        value=len(st.session_state.evaluation_criteria) if st.session_state.evaluation_criteria else 4,
        step=1,
        key="num_criteria"
    )
    
    # 평가 기준 입력 폼
    criteria_list = []
    
    for i in range(num_criteria):
        with st.expander(f"평가 기준 {i+1}", expanded=True):
            # 저장된 평가 기준이 있으면 해당 값 사용
            saved_criterion = None
            if i < len(st.session_state.evaluation_criteria):
                saved_criterion = st.session_state.evaluation_criteria[i]
            
            # 기준명 입력
            criterion_name = st.text_input(
                f"기준명 {i+1}",
                placeholder="예: 내용의 충실성, 논리성 등",
                key=f"criterion_name_{i}",
                value=saved_criterion["name"] if saved_criterion else ""
            )
            
            # 기준 상세 설명 입력
            criterion_description = st.text_area(
                f"기준 상세 설명 {i+1}",
                placeholder="이 평가 기준의 구체적인 평가 내용을 설명하세요.",
                key=f"criterion_description_{i}",
                value=saved_criterion.get("description", "") if saved_criterion else "",
                height=80
            )
            
            # 점수 입력
            col1, col2 = st.columns(2)
            with col1:
                max_score = st.number_input(
                    "최고점",
                    min_value=0.0,
                    max_value=100.0,
                    value=saved_criterion["max_score"] if saved_criterion else 25.0,
                    step=0.5,
                    key=f"max_score_{i}"
                )
            
            with col2:
                min_score = st.number_input(
                    "최저점",
                    min_value=0.0,
                    max_value=100.0,
                    value=saved_criterion["min_score"] if saved_criterion else 15.0,
                    step=0.5,
                    key=f"min_score_{i}"
                )
            
            # 가중치 입력
            st.markdown("**가중치 설정**")
            weight = st.number_input(
                f"가중치 {i+1}",
                min_value=0.0,
                max_value=10.0,
                value=saved_criterion.get("weight", 1.0) if saved_criterion else 1.0,
                step=0.1,
                key=f"weight_{i}",
                help="총점 계산 시 이 평가 기준의 중요도를 나타냅니다. 기본값은 1.0입니다. 예: 1.5는 1.5배 가중치를 의미합니다."
            )
            
            if criterion_name:
                criteria_list.append({
                    "name": criterion_name,
                    "description": criterion_description,
                    "max_score": max_score,
                    "min_score": min_score,
                    "weight": weight
                })
    
    # 평가 기준 저장
    if criteria_list:
        st.session_state.evaluation_criteria = criteria_list
        
        # 평가 기준 저장 버튼
        st.markdown("---")
        st.markdown("### 💾 평가 기준 저장")
        
        # session_state의 evaluation_title 사용 (확인 버튼으로 설정된 값)
        if st.session_state.evaluation_title:
            save_key = f"save_criteria_{st.session_state.evaluation_title}"
            
            col1, col2 = st.columns([2, 1])
            with col1:
                st.info(f"💡 현재 평가 기준을 '{st.session_state.evaluation_title}' 제목으로 저장하시겠습니까?")
            
            with col2:
                if st.button("💾 평가 기준 저장", key=save_key, use_container_width=True, type="primary"):
                    # 평가 기준을 딕셔너리 형태로 저장 (깊은 복사)
                    import copy
                    st.session_state.saved_criteria_templates[st.session_state.evaluation_title] = copy.deepcopy(criteria_list)
                    # 파일에 저장
                    save_criteria_templates(st.session_state.saved_criteria_templates)
                    st.success(f"✅ '{st.session_state.evaluation_title}' 평가 기준이 저장되었습니다!")
                    st.rerun()
            
            # 이미 저장된 제목인지 확인
            if st.session_state.evaluation_title in st.session_state.saved_criteria_templates:
                st.warning(f"⚠️ '{st.session_state.evaluation_title}' 제목으로 이미 저장된 평가 기준이 있습니다. 저장하면 기존 내용이 덮어씌워집니다.")
        else:
            st.info("💡 평가 제목을 먼저 입력하고 확인 버튼을 눌러주시면 평가 기준을 저장할 수 있습니다.")
    
    st.markdown("---")
    
    # 3. PDF 파일 업로드
    st.header("3️⃣ PDF 파일 업로드")
    
    uploaded_files = st.file_uploader(
        "에세이 PDF 파일을 업로드하세요",
        type=['pdf'],
        accept_multiple_files=True,
        help="여러 개의 PDF 파일을 동시에 업로드할 수 있습니다."
    )
    
    if uploaded_files:
        st.session_state.uploaded_pdfs = uploaded_files
        
        # PDF 텍스트 추출
        if st.button("📄 PDF 텍스트 추출하기", type="primary"):
            st.session_state.extracted_texts = []
            
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            for idx, pdf_file in enumerate(uploaded_files):
                status_text.text(f"처리 중: {pdf_file.name} ({idx+1}/{len(uploaded_files)})")
                
                text = extract_text_from_pdf(pdf_file)
                
                st.session_state.extracted_texts.append({
                    "filename": pdf_file.name,
                    "text": text
                })
                
                progress_bar.progress((idx + 1) / len(uploaded_files))
            
            status_text.text("✅ 모든 PDF 파일 처리가 완료되었습니다!")
            progress_bar.empty()
            st.success(f"{len(uploaded_files)}개의 PDF 파일에서 텍스트를 추출했습니다.")
            
            # 텍스트 추출 완료 후 업로드된 PDF 리스트 삭제
            st.session_state.uploaded_pdfs = []
            st.rerun()
            
            # 텍스트 추출 완료 후 업로드된 PDF 리스트 삭제
            st.session_state.uploaded_pdfs = []
    
    st.markdown("---")
    
    # 4. 추출된 텍스트 미리보기
    if st.session_state.extracted_texts:
        st.header("4️⃣ 추출된 텍스트 미리보기")
        
        for idx, extracted in enumerate(st.session_state.extracted_texts):
            with st.expander(f"📄 {extracted['filename']}", expanded=False):
                if extracted['text']:
                    st.text_area(
                        "추출된 텍스트",
                        extracted['text'],
                        height=200,
                        key=f"preview_{idx}",
                        disabled=True
                    )
                else:
                    st.warning("이 PDF에서 텍스트를 추출할 수 없습니다.")
        
        st.markdown("---")
        
        # 5. 평가하기 버튼
        st.header("5️⃣ 평가 실행")
        
        if st.button("🔍 평가하기", type="primary", use_container_width=True):
            # 유효성 검사
            if not st.session_state.evaluation_criteria:
                st.error("⚠️ 평가 기준을 먼저 설정해주세요!")
            elif not OPENAI_API_KEY:
                st.error("⚠️ OpenAI API Key가 설정되지 않았습니다! .env 파일에 OPENAI_API_KEY를 설정해주세요.")
            else:
                # 평가 결과 초기화
                st.session_state.evaluation_results = []
                
                # 진행 상황 표시
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                # 각 학생(PDF)별로 평가 결과 생성
                for idx, extracted in enumerate(st.session_state.extracted_texts):
                    status_text.text(f"평가 중: {extracted['filename']} ({idx+1}/{len(st.session_state.extracted_texts)})")
                    
                    # 표절 검사를 포함한 평가 수행
                    evaluation_result = evaluate_essay_with_plagiarism_check(
                        extracted['text'],
                        extracted['filename'],
                        st.session_state.evaluation_criteria,
                        OPENAI_API_KEY,
                        st.session_state.evaluated_essays
                    )
                    
                    if evaluation_result:
                        result = {
                            "filename": extracted['filename'],
                            "scores": evaluation_result["scores"],
                            "total_score": evaluation_result["total_score"],
                            "feedback": evaluation_result["feedback"]
                        }
                        # 표절 검사 정보가 있으면 추가
                        if 'plagiarism_check' in evaluation_result:
                            result['plagiarism_check'] = evaluation_result['plagiarism_check']
                        
                        st.session_state.evaluation_results.append(result)
                        
                        # 평가 완료된 에세이를 저장 (표절 검사용)
                        st.session_state.evaluated_essays.append({
                            "filename": extracted['filename'],
                            "text": extracted['text']
                        })
                    else:
                        # 오류 발생 시 기본값
                        result = {
                            "filename": extracted['filename'],
                            "scores": {criterion["name"]: 0.0 for criterion in st.session_state.evaluation_criteria},
                            "total_score": 0.0,
                            "feedback": "평가 중 오류가 발생했습니다."
                        }
                        st.session_state.evaluation_results.append(result)
                    
                    progress_bar.progress((idx + 1) / len(st.session_state.extracted_texts))
                
                status_text.text("✅ 모든 평가가 완료되었습니다!")
                progress_bar.empty()
                st.success(f"✅ {len(st.session_state.extracted_texts)}개의 에세이 평가가 완료되었습니다!")
                st.rerun()
        
        st.markdown("---")
        
        # 6. 평가 결과 표시
        if st.session_state.evaluation_results:
            st.header("6️⃣ 평가 결과")
            
            # 평가 정보 표시
            info_cols = st.columns(4)
            with info_cols[0]:
                if st.session_state.evaluation_year:
                    st.metric("📅 평가 년도", st.session_state.evaluation_year)
            with info_cols[1]:
                if st.session_state.evaluation_semester:
                    st.metric("📚 학기", st.session_state.evaluation_semester)
            with info_cols[2]:
                if st.session_state.evaluation_subject:
                    st.metric("📖 과목명", st.session_state.evaluation_subject)
            with info_cols[3]:
                if st.session_state.evaluation_title:
                    st.metric("📌 평가 제목", st.session_state.evaluation_title[:20] + "..." if len(st.session_state.evaluation_title) > 20 else st.session_state.evaluation_title)
            
            if st.session_state.evaluation_title:
                st.markdown(f"### 📌 {st.session_state.evaluation_title}")
            
            # 결과 요약 테이블
            st.subheader("📊 전체 학생 점수 요약")
            
            # 총점 최고점 계산 (가중치 반영)
            total_max_score = sum(c["max_score"] * c.get("weight", 1.0) for c in st.session_state.evaluation_criteria)
            
            # 만점 조정 기능
            st.markdown("### ⚙️ 만점 조정")
            col1, col2, col3 = st.columns([2, 2, 1])
            with col1:
                adjusted_max = st.number_input(
                    "만점 점수 설정",
                    min_value=0.0,
                    max_value=1000.0,
                    value=float(st.session_state.adjusted_max_score) if st.session_state.adjusted_max_score else total_max_score,
                    step=1.0,
                    key="adjusted_max_score_input",
                    help=f"현재 총점 만점: {total_max_score:.1f}점. 다른 만점으로 조정할 수 있습니다."
                )
            
            with col2:
                if st.button("✅ 만점 적용", key="apply_adjusted_max", use_container_width=True):
                    st.session_state.adjusted_max_score = adjusted_max
                    st.success(f"✅ 만점이 {adjusted_max:.1f}점으로 설정되었습니다!")
                    st.rerun()
            
            with col3:
                if st.button("🔄 원래 점수로", key="reset_adjusted_max", use_container_width=True):
                    st.session_state.adjusted_max_score = None
                    st.success("✅ 원래 점수로 되돌렸습니다!")
                    st.rerun()
            
            # 만점이 조정되었는지 확인
            use_adjusted = st.session_state.adjusted_max_score is not None
            target_max = float(st.session_state.adjusted_max_score) if use_adjusted else total_max_score
            
            if use_adjusted:
                st.info(f"💡 점수가 {target_max:.1f}점 만점으로 조정되어 표시됩니다. (원래 만점: {total_max_score:.1f}점)")
            
            # 테이블 데이터 준비 (원래 점수와 조정된 총점 모두 표시)
            summary_data = {
                "학생": [],
                **{criterion["name"]: [] for criterion in st.session_state.evaluation_criteria},
                "총점(원래)": []
            }
            
            # 만점 조정이 적용된 경우 조정된 총점 열 추가
            if use_adjusted:
                summary_data["총점(조정)"] = []
            
            for result in st.session_state.evaluation_results:
                # 파일명에서 확장자 제거하여 학생명으로 표시
                student_name = result["filename"].replace(".pdf", "").replace(".PDF", "")
                summary_data["학생"].append(student_name)
                
                # 각 기준별 점수 표시
                for criterion in st.session_state.evaluation_criteria:
                    original_score = result["scores"].get(criterion["name"], 0.0)
                    # 만점 조정이 적용된 경우 조정된 점수와 원래 점수를 함께 표시
                    if use_adjusted:
                        criterion_max = criterion["max_score"] * criterion.get("weight", 1.0)
                        adjusted_criterion_max = (criterion_max / total_max_score) * target_max
                        adjusted_score = (original_score / criterion_max) * adjusted_criterion_max if criterion_max > 0 else 0
                        summary_data[criterion["name"]].append(f"{adjusted_score:.1f} ({original_score:.1f})")
                    else:
                        summary_data[criterion["name"]].append(f"{original_score:.1f}")
                
                # 원래 총점 표시
                original_total = result['total_score']
                summary_data["총점(원래)"].append(f"{original_total:.1f}")
                
                # 조정된 총점 표시 (만점 조정이 적용된 경우만, 원래 점수를 괄호 안에 표시)
                if use_adjusted:
                    adjusted_total = (result['total_score'] / total_max_score) * target_max if total_max_score > 0 else 0
                    summary_data["총점(조정)"].append(f"{adjusted_total:.1f} ({original_total:.1f})")
            
            # pandas DataFrame으로 변환
            df = pd.DataFrame(summary_data)
            
            # 표에 최고점 및 가중치 정보 추가 표시
            st.markdown("**평가 기준별 최고점 및 가중치:**")
            criteria_info_list = []
            for c in st.session_state.evaluation_criteria:
                weight = c.get("weight", 1.0)
                criterion_max = c["max_score"] * weight
                if use_adjusted:
                    adjusted_criterion_max = (criterion_max / total_max_score) * target_max
                    if weight != 1.0:
                        criteria_info_list.append(f"{c['name']}: {adjusted_criterion_max:.1f}점 (가중치: {weight})")
                    else:
                        criteria_info_list.append(f"{c['name']}: {adjusted_criterion_max:.1f}점")
                else:
                    if weight != 1.0:
                        criteria_info_list.append(f"{c['name']}: {c['max_score']}점 (가중치: {weight})")
                    else:
                        criteria_info_list.append(f"{c['name']}: {c['max_score']}점")
            criteria_info = " | ".join(criteria_info_list)
            if use_adjusted:
                st.markdown(f"*{criteria_info} | 총점(원래): {total_max_score:.1f}점 | 총점(조정): {target_max:.1f}점*")
            else:
                st.markdown(f"*{criteria_info} | 총점: {total_max_score:.1f}점*")
            st.markdown("")
            
            # 데이터프레임 표시
            st.dataframe(df, use_container_width=True, hide_index=True)
            
            # 엑셀 다운로드 버튼
            st.markdown("---")
            st.subheader("📥 결과 다운로드")
            
            # 엑셀 파일 생성 함수 (원래 점수)
            def create_excel_file(use_adjusted_scores=False):
                # Excel 파일을 위한 데이터 준비 (피드백 제외, 점수만 포함)
                excel_data = {
                    "학생": [],
                    **{criterion["name"]: [] for criterion in st.session_state.evaluation_criteria},
                    "총점": []
                }
                
                for result in st.session_state.evaluation_results:
                    student_name = result["filename"].replace(".pdf", "").replace(".PDF", "")
                    excel_data["학생"].append(student_name)
                    
                    for criterion in st.session_state.evaluation_criteria:
                        score = result["scores"].get(criterion["name"], 0.0)
                        # 만점 조정이 적용된 경우 점수 비율 조정
                        if use_adjusted_scores and use_adjusted:
                            criterion_max = criterion["max_score"] * criterion.get("weight", 1.0)
                            adjusted_criterion_max = (criterion_max / total_max_score) * target_max
                            adjusted_score = (score / criterion_max) * adjusted_criterion_max if criterion_max > 0 else 0
                            # 소수점 첫째자리까지 반올림
                            excel_data[criterion["name"]].append(round(adjusted_score, 1))
                        else:
                            excel_data[criterion["name"]].append(score)
                    
                    # 총점 조정
                    if use_adjusted_scores and use_adjusted:
                        adjusted_total = (result['total_score'] / total_max_score) * target_max if total_max_score > 0 else 0
                        # 소수점 첫째자리까지 반올림
                        excel_data["총점"].append(round(adjusted_total, 1))
                    else:
                        excel_data["총점"].append(result['total_score'])
                
                # DataFrame 생성
                excel_df = pd.DataFrame(excel_data)
                
                # Excel 파일을 메모리에 생성
                output = BytesIO()
                with pd.ExcelWriter(output, engine='openpyxl') as writer:
                    # 점수 요약 시트
                    excel_df.to_excel(writer, sheet_name='점수 요약', index=False)
                    
                    # 평가 기준 정보 시트
                    criteria_data = {
                        "평가 기준": [c["name"] for c in st.session_state.evaluation_criteria],
                        "기준 상세 설명": [c.get("description", "") for c in st.session_state.evaluation_criteria],
                        "최저점": [c["min_score"] for c in st.session_state.evaluation_criteria],
                        "최고점": [c["max_score"] for c in st.session_state.evaluation_criteria],
                        "가중치": [c.get("weight", 1.0) for c in st.session_state.evaluation_criteria]
                    }
                    
                    # 만점 조정이 적용된 경우 조정된 최고점도 표시
                    if use_adjusted_scores and use_adjusted:
                        criteria_data["조정된 최고점"] = [
                            (c["max_score"] * c.get("weight", 1.0) / total_max_score) * target_max 
                            for c in st.session_state.evaluation_criteria
                        ]
                    
                    criteria_df = pd.DataFrame(criteria_data)
                    criteria_df.to_excel(writer, sheet_name='평가 기준', index=False)
                    
                    # 평가 정보 시트
                    info_data = {
                        "항목": ["평가 년도", "학기", "과목명", "평가 제목", "원래 총점 만점", "조정된 총점 만점"],
                        "내용": [
                            st.session_state.evaluation_year or "",
                            st.session_state.evaluation_semester or "",
                            st.session_state.evaluation_subject or "",
                            st.session_state.evaluation_title or "",
                            f"{total_max_score:.1f}점",
                            f"{target_max:.1f}점" if use_adjusted_scores and use_adjusted else f"{total_max_score:.1f}점"
                        ]
                    }
                    info_df = pd.DataFrame(info_data)
                    info_df.to_excel(writer, sheet_name='평가 정보', index=False)
                
                output.seek(0)
                return output.getvalue()
            
            # 점수 누적 기능 (동일 제목으로 평가할 때만 누적)
            st.markdown("### 💾 점수 누적 저장")
            
            # 평가 제목이 없으면 누적 기능 비활성화
            if not st.session_state.evaluation_title:
                st.warning("⚠️ 평가 제목을 설정해야 누적 저장 기능을 사용할 수 있습니다.")
                st.markdown("---")
            else:
                # 누적 파일 경로 (평가 제목별로 파일 생성)
                accumulated_file = f"누적점수_{st.session_state.evaluation_title}.xlsx"
            
            def load_accumulated_data():
                """기존 누적 데이터를 로드합니다."""
                if os.path.exists(accumulated_file):
                    try:
                        df = pd.read_excel(accumulated_file, sheet_name='점수 요약')
                        return df
                    except Exception as e:
                        st.warning(f"기존 누적 파일을 읽는 중 오류 발생: {str(e)}")
                        return pd.DataFrame()
                return pd.DataFrame()
            
            def save_accumulated_data(new_data_df):
                """새로운 데이터를 기존 데이터에 추가하여 저장합니다."""
                # 기존 데이터 로드
                existing_df = load_accumulated_data()
                
                if existing_df.empty:
                    # 기존 데이터가 없으면 새로 생성
                    combined_df = new_data_df.copy()
                else:
                    # 기존 데이터와 새 데이터 병합
                    # 학생명을 기준으로 중복 제거 (새 데이터가 우선)
                    # 기존 데이터에서 새 데이터에 있는 학생 제거
                    existing_students = set(existing_df['학생'].values) if '학생' in existing_df.columns else set()
                    new_students = set(new_data_df['학생'].values) if '학생' in new_data_df.columns else set()
                    
                    # 기존 데이터에서 새 데이터에 포함된 학생 제거
                    existing_df_filtered = existing_df[~existing_df['학생'].isin(new_students)]
                    
                    # 기존 데이터와 새 데이터 결합
                    combined_df = pd.concat([existing_df_filtered, new_data_df], ignore_index=True)
                
                # 평가 정보도 함께 저장
                try:
                    with pd.ExcelWriter(accumulated_file, engine='openpyxl') as writer:
                        # 점수 요약 시트
                        combined_df.to_excel(writer, sheet_name='점수 요약', index=False)
                        
                        # 평가 기준 정보 시트 (현재 기준으로 업데이트)
                        criteria_data = {
                            "평가 기준": [c["name"] for c in st.session_state.evaluation_criteria],
                            "기준 상세 설명": [c.get("description", "") for c in st.session_state.evaluation_criteria],
                            "최저점": [c["min_score"] for c in st.session_state.evaluation_criteria],
                            "최고점": [c["max_score"] for c in st.session_state.evaluation_criteria],
                            "가중치": [c.get("weight", 1.0) for c in st.session_state.evaluation_criteria]
                        }
                        criteria_df = pd.DataFrame(criteria_data)
                        criteria_df.to_excel(writer, sheet_name='평가 기준', index=False)
                        
                        # 평가 정보 시트
                        info_data = {
                            "항목": ["평가 년도", "학기", "과목명", "평가 제목", "총점 만점", "마지막 업데이트"],
                            "내용": [
                                st.session_state.evaluation_year or "",
                                st.session_state.evaluation_semester or "",
                                st.session_state.evaluation_subject or "",
                                st.session_state.evaluation_title or "",
                                f"{total_max_score:.1f}점",
                                pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S")
                            ]
                        }
                        info_df = pd.DataFrame(info_data)
                        info_df.to_excel(writer, sheet_name='평가 정보', index=False)
                    
                    return True, combined_df
                except Exception as e:
                    return False, str(e)
            
            # 현재 결과를 DataFrame으로 변환
            current_data = {
                "학생": [],
                **{criterion["name"]: [] for criterion in st.session_state.evaluation_criteria},
                "총점": []
            }
            
            for result in st.session_state.evaluation_results:
                student_name = result["filename"].replace(".pdf", "").replace(".PDF", "")
                current_data["학생"].append(student_name)
                
                for criterion in st.session_state.evaluation_criteria:
                    score = result["scores"].get(criterion["name"], 0.0)
                    current_data[criterion["name"]].append(score)
                
                current_data["총점"].append(result['total_score'])
            
            current_df = pd.DataFrame(current_data)
            
            # 엑셀 파일 저장
            col1, col2, col3 = st.columns([1, 1, 1])
            
            with col1:
                if st.button("💾 엑셀에 누적 저장", use_container_width=True, type="primary"):
                    if not current_df.empty:
                        # 저장 전 확인 메시지
                        st.markdown("---")
                        st.markdown("### ⚠️ 저장 전 확인")
                        st.info(f"""
                        **저장 정보 확인:**
                        - 📌 평가 제목: {st.session_state.evaluation_title}
                        - 📁 파일명: {accumulated_file}
                        - 👥 저장할 학생 수: {len(current_df)}명
                        - 📅 평가 년도: {st.session_state.evaluation_year or 'N/A'}
                        - 📚 학기: {st.session_state.evaluation_semester or 'N/A'}
                        - 📖 과목명: {st.session_state.evaluation_subject or 'N/A'}
                        
                        **저장할 학생 목록:**
                        {', '.join(current_df['학생'].tolist())}
                        """)
                        
                        # 확인 버튼
                        col_confirm1, col_confirm2 = st.columns(2)
                        with col_confirm1:
                            if st.button("✅ 확인하고 저장", use_container_width=True, type="primary", key="confirm_save"):
                                success, result = save_accumulated_data(current_df)
                                if success:
                                    st.success(f"✅ {len(current_df)}명의 점수가 누적 파일에 저장되었습니다!")
                                    st.info(f"📁 파일명: {accumulated_file}")
                                    st.rerun()
                                else:
                                    st.error(f"❌ 저장 중 오류 발생: {result}")
                        with col_confirm2:
                            if st.button("❌ 취소", use_container_width=True, key="cancel_save"):
                                st.rerun()
                    else:
                        st.warning("⚠️ 저장할 데이터가 없습니다.")
            
            with col2:
                # 누적 데이터 보기
                accumulated_df = load_accumulated_data()
                if not accumulated_df.empty:
                    st.info(f"📊 누적된 학생 수: {len(accumulated_df)}명")
                    if st.button("📋 누적 데이터 보기", use_container_width=True):
                        st.session_state.show_accumulated = not st.session_state.show_accumulated
                        st.rerun()
                else:
                    st.info("💡 아직 누적된 데이터가 없습니다.")
            
            with col3:
                # 누적 파일 다운로드
                if os.path.exists(accumulated_file):
                    with open(accumulated_file, 'rb') as f:
                        st.download_button(
                            label="📥 누적 파일 다운로드",
                            data=f.read(),
                            file_name=accumulated_file,
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            use_container_width=True
                        )
            
            # 누적 데이터 표시
            if st.session_state.get('show_accumulated', False):
                accumulated_df = load_accumulated_data()
                if not accumulated_df.empty:
                    st.markdown("---")
                    st.subheader("📊 누적된 전체 점수 데이터")
                    st.dataframe(accumulated_df, use_container_width=True, hide_index=True)
                    st.caption(f"총 {len(accumulated_df)}명의 학생 데이터가 누적되어 있습니다.")
            
            st.markdown("---")
            
            # 다운로드 버튼들
            st.markdown("### 📥 현재 결과 다운로드")
            col1, col2 = st.columns(2)
            
            with col1:
                # 원래 점수 다운로드
                excel_file_original = create_excel_file(use_adjusted_scores=False)
                filename_original = f"에세이평가결과_원래점수_{st.session_state.evaluation_year or 'N/A'}_{st.session_state.evaluation_semester or 'N/A'}.xlsx"
                
                st.download_button(
                    label="📥 원래 점수로 다운로드",
                    data=excel_file_original,
                    file_name=filename_original,
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True,
                    type="primary"
                )
            
            with col2:
                # 조정된 점수 다운로드 (만점 조정이 적용된 경우만 활성화)
                if use_adjusted:
                    excel_file_adjusted = create_excel_file(use_adjusted_scores=True)
                    filename_adjusted = f"에세이평가결과_조정점수({target_max:.0f}점만점)_{st.session_state.evaluation_year or 'N/A'}_{st.session_state.evaluation_semester or 'N/A'}.xlsx"
                    
                    st.download_button(
                        label=f"📥 조정된 점수로 다운로드 ({target_max:.0f}점 만점)",
                        data=excel_file_adjusted,
                        file_name=filename_adjusted,
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        use_container_width=True,
                        type="primary"
                    )
                else:
                    st.info("💡 만점을 조정하면 조정된 점수로 다운로드할 수 있습니다.")
            
            st.markdown("---")
            
            # 학생별 상세 피드백
            st.subheader("📝 학생별 상세 피드백")
            
            # 일괄 다운로드 버튼 (상단에 배치)
            st.markdown("### 📥 피드백 보고서 다운로드")
            
            def create_all_reports_zip():
                """모든 학생의 피드백 보고서를 ZIP 파일로 생성합니다."""
                zip_buffer = BytesIO()
                
                evaluation_info = {
                    'year': st.session_state.evaluation_year,
                    'semester': st.session_state.evaluation_semester,
                    'subject': st.session_state.evaluation_subject,
                    'title': st.session_state.evaluation_title
                }
                
                with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                    for result in st.session_state.evaluation_results:
                        student_name = result['filename'].replace('.pdf', '').replace('.PDF', '')
                        report = create_feedback_report(result, st.session_state.evaluation_criteria, evaluation_info)
                        zip_file.writestr(f"{student_name}_피드백보고서.docx", report.getvalue())
                
                zip_buffer.seek(0)
                return zip_buffer.getvalue()
            
            # 일괄 다운로드 버튼
            all_reports_zip = create_all_reports_zip()
            zip_filename = f"전체_피드백보고서_{st.session_state.evaluation_year or 'N/A'}_{st.session_state.evaluation_semester or 'N/A'}.zip"
            
            st.info(f"💡 전체 {len(st.session_state.evaluation_results)}명의 피드백 보고서를 한 번에 다운로드할 수 있습니다.")
            
            st.download_button(
                label=f"📦 전체 피드백 보고서 일괄 다운로드 (ZIP) - {len(st.session_state.evaluation_results)}개 파일",
                data=all_reports_zip,
                file_name=zip_filename,
                mime="application/zip",
                use_container_width=True,
                type="primary"
            )
            
            st.markdown("---")
            st.markdown("### 👤 개별 피드백 보고서")
            
            for result in st.session_state.evaluation_results:
                # 파일명에서 확장자 제거
                student_name = result["filename"].replace(".pdf", "").replace(".PDF", "")
                
                with st.expander(f"👤 {student_name} ({result['filename']})", expanded=False):
                    # 항목별 점수 카드
                    st.markdown("### 📊 항목별 점수")
                    
                    # 점수 표시를 더 예쁘게
                    cols = st.columns(len(st.session_state.evaluation_criteria) + 1)
                    for idx, criterion in enumerate(st.session_state.evaluation_criteria):
                        score = result["scores"].get(criterion["name"], 0.0)
                        max_score = criterion["max_score"]
                        percentage = (score / max_score * 100) if max_score > 0 else 0
                        
                        with cols[idx]:
                            # 점수에 따라 색상 구분
                            if percentage >= 80:
                                delta_color = "normal"
                            elif percentage >= 60:
                                delta_color = "normal"
                            else:
                                delta_color = "inverse"
                            
                            st.metric(
                                criterion["name"],
                                f"{score:.1f}",
                                delta=f"/ {max_score:.1f}점",
                                delta_color=delta_color
                            )
                    
                    with cols[-1]:
                        # 총점 계산 (가중치 반영)
                        total_max = sum(c["max_score"] * c.get("weight", 1.0) for c in st.session_state.evaluation_criteria)
                        total_percentage = (result['total_score'] / total_max * 100) if total_max > 0 else 0
                        
                        st.metric(
                            "총점",
                            f"{result['total_score']:.1f}",
                            delta=f"/ {total_max:.1f}점",
                            delta_color="normal"
                        )
                    
                    st.markdown("---")
                    
                    # 상세 피드백
                    st.markdown("### 💬 상세 피드백")
                    st.markdown("---")
                    
                    # 피드백을 파싱하여 구조화된 데이터로 변환
                    parsed_feedback = parse_feedback(result["feedback"], st.session_state.evaluation_criteria)
                    
                    # 각 평가 기준별 피드백을 깔끔하게 표시
                    for criterion in st.session_state.evaluation_criteria:
                        criterion_name = criterion['name']
                        item_feedback = parsed_feedback['items'].get(criterion_name, {})
                        summary = item_feedback.get('summary', '')
                        good_points = item_feedback.get('good_points', '')
                        improvement_points = item_feedback.get('improvement_points', '')
                        
                        # 평가 기준명 표시
                        st.markdown(f"#### 📌 {criterion_name}")
                        
                        # 평가 요약
                        if summary:
                            st.markdown("**【평가 요약】**")
                            st.markdown(summary)
                            st.markdown("")
                        
                        # 잘 작성한 점
                        if good_points:
                            st.markdown("**✨ 잘 작성한 점:**")
                            st.markdown(good_points)
                            st.markdown("")
                        
                        # 개선할 점 및 오류
                        if improvement_points:
                            st.markdown("**⚠️ 개선할 점 및 오류:**")
                            st.markdown(improvement_points)
                            st.markdown("")
                        
                        # 피드백이 없는 경우
                        if not summary and not good_points and not improvement_points:
                            st.info("피드백 없음")
                        
                        st.markdown("---")
                    
                    # 종합의견 추가
                    if parsed_feedback['general']:
                        st.markdown("#### 📌 **종합의견**")
                        st.markdown(f"**{parsed_feedback['general']}**")
                        st.markdown("---")
                    
                    # 개별 다운로드 버튼
                    st.markdown("---")
                    st.markdown("#### 📄 개별 보고서 다운로드")
                    
                    evaluation_info = {
                        'year': st.session_state.evaluation_year,
                        'semester': st.session_state.evaluation_semester,
                        'subject': st.session_state.evaluation_subject,
                        'title': st.session_state.evaluation_title
                    }
                    
                    report_file = create_feedback_report(result, st.session_state.evaluation_criteria, evaluation_info)
                    report_filename = f"{student_name}_피드백보고서.docx"
                    
                    st.download_button(
                        label=f"📥 {student_name} 피드백 보고서 다운로드",
                        data=report_file.getvalue(),
                        file_name=report_filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"download_{result['filename']}",
                        use_container_width=True
                    )
            
            st.markdown("---")
            
            # 7. 누적 엑셀 파일 분석 및 시각화
            st.header("7️⃣ 누적 점수 분석 및 시각화")
            
            uploaded_analysis_file = st.file_uploader(
                "누적 점수 엑셀 파일을 업로드하여 점수 분포를 분석하세요",
                type=['xlsx', 'xls'],
                help="누적 점수 엑셀 파일을 업로드하면 점수 분포 히스토그램과 하위 20% 학생 리스트를 확인할 수 있습니다."
            )
            
            if uploaded_analysis_file:
                try:
                    # 엑셀 파일 읽기
                    df_analysis = pd.read_excel(uploaded_analysis_file, sheet_name='점수 요약')
                    
                    # 총점 열 찾기
                    total_score_column = None
                    for col in df_analysis.columns:
                        if '총점' in col:
                            total_score_column = col
                            break
                    
                    if total_score_column is None:
                        st.error("⚠️ 엑셀 파일에서 '총점' 열을 찾을 수 없습니다.")
                    else:
                        # 총점 데이터 추출
                        scores = df_analysis[total_score_column].dropna().astype(float)
                        
                        if len(scores) == 0:
                            st.error("⚠️ 분석할 점수 데이터가 없습니다.")
                        else:
                            # 하위 20% 계산
                            percentile_20 = np.percentile(scores, 20)
                            bottom_20_percent = df_analysis[scores <= percentile_20].copy()
                            
                            # 통계 정보 계산
                            mean_score = scores.mean()
                            median_score = scores.median()
                            std_score = scores.std()
                            
                            # seaborn 스타일 설정
                            sns.set_style("whitegrid")
                            sns.set_palette("husl")
                            
                            # 탭으로 여러 시각화 제공
                            tab1, tab2, tab3 = st.tabs(["📊 분포도 (히스토그램 + 밀도)", "📦 박스플롯", "📈 통합 분석"])
                            
                            with tab1:
                                # 히스토그램 + KDE 밀도 곡선
                                fig, ax = plt.subplots(figsize=(12, 7))
                                
                                # 히스토그램과 KDE를 함께 표시
                                sns.histplot(
                                    scores, 
                                    bins=25, 
                                    kde=True, 
                                    color='#2E86AB', 
                                    alpha=0.7,
                                    edgecolor='white',
                                    linewidth=1.5,
                                    kde_kws={'linewidth': 3, 'color': '#A23B72'}
                                )
                                
                                # 하위 20% 영역 강조
                                ax.axvspan(0, percentile_20, alpha=0.2, color='#F18F01', label=f'하위 20% 영역 (≤{percentile_20:.1f}점)')
                                
                                # 통계선 표시
                                ax.axvline(mean_score, color='#06A77D', linestyle='--', linewidth=2.5, alpha=0.9, label=f'평균: {mean_score:.1f}점')
                                ax.axvline(median_score, color='#D56062', linestyle='--', linewidth=2.5, alpha=0.9, label=f'중앙값: {median_score:.1f}점')
                                ax.axvline(percentile_20, color='#F18F01', linestyle='-', linewidth=2, alpha=0.8, label=f'하위 20% 경계: {percentile_20:.1f}점')
                                
                                # 그래프 스타일링
                                ax.set_xlabel('총점', fontsize=13, fontweight='bold', color='#2C3E50')
                                ax.set_ylabel('학생 수', fontsize=13, fontweight='bold', color='#2C3E50')
                                ax.set_title('학생 점수 분포 분석', fontsize=16, fontweight='bold', pad=20, color='#2C3E50')
                                ax.legend(loc='upper right', fontsize=10, framealpha=0.9)
                                ax.grid(True, alpha=0.3, linestyle='--')
                                
                                # 배경색 설정
                                ax.set_facecolor('#F8F9FA')
                                fig.patch.set_facecolor('white')
                                
                                plt.tight_layout()
                                st.pyplot(fig)
                                plt.close(fig)
                            
                            with tab2:
                                # 박스플롯 + 바이올린 플롯
                                fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(14, 6))
                                
                                # 박스플롯
                                bp = ax1.boxplot(
                                    [scores], 
                                    vert=True,
                                    patch_artist=True,
                                    boxprops=dict(facecolor='#2E86AB', alpha=0.7),
                                    medianprops=dict(color='#D56062', linewidth=2.5),
                                    whiskerprops=dict(color='#2C3E50', linewidth=1.5),
                                    capprops=dict(color='#2C3E50', linewidth=1.5)
                                )
                                
                                # 하위 20% 경계선
                                ax1.axhline(percentile_20, color='#F18F01', linestyle='--', linewidth=2, label=f'하위 20% 경계: {percentile_20:.1f}점')
                                ax1.axhline(mean_score, color='#06A77D', linestyle='--', linewidth=2, label=f'평균: {mean_score:.1f}점')
                                
                                ax1.set_ylabel('총점', fontsize=12, fontweight='bold', color='#2C3E50')
                                ax1.set_title('박스플롯 (분포 요약)', fontsize=14, fontweight='bold', color='#2C3E50')
                                ax1.grid(True, alpha=0.3, linestyle='--')
                                ax1.legend(fontsize=9)
                                ax1.set_facecolor('#F8F9FA')
                                
                                # 바이올린 플롯
                                parts = ax2.violinplot(
                                    [scores],
                                    positions=[1],
                                    showmeans=True,
                                    showmedians=True,
                                    widths=0.6
                                )
                                
                                # 바이올린 플롯 색상 설정
                                for pc in parts['bodies']:
                                    pc.set_facecolor('#2E86AB')
                                    pc.set_alpha(0.7)
                                
                                parts['cmeans'].set_color('#06A77D')
                                parts['cmeans'].set_linewidth(2)
                                parts['cmedians'].set_color('#D56062')
                                parts['cmedians'].set_linewidth(2)
                                
                                ax2.axhline(percentile_20, color='#F18F01', linestyle='--', linewidth=2, label=f'하위 20% 경계: {percentile_20:.1f}점')
                                ax2.set_ylabel('총점', fontsize=12, fontweight='bold', color='#2C3E50')
                                ax2.set_title('바이올린 플롯 (밀도 분포)', fontsize=14, fontweight='bold', color='#2C3E50')
                                ax2.set_xticks([1])
                                ax2.set_xticklabels(['점수 분포'])
                                ax2.grid(True, alpha=0.3, linestyle='--')
                                ax2.legend(fontsize=9)
                                ax2.set_facecolor('#F8F9FA')
                                
                                fig.patch.set_facecolor('white')
                                plt.tight_layout()
                                st.pyplot(fig)
                                plt.close(fig)
                            
                            with tab3:
                                # 통합 분석: 히스토그램 + 박스플롯 + 통계
                                fig = plt.figure(figsize=(14, 8))
                                gs = fig.add_gridspec(3, 2, hspace=0.3, wspace=0.3)
                                
                                # 메인 히스토그램 (상단 전체)
                                ax_main = fig.add_subplot(gs[0:2, :])
                                sns.histplot(
                                    scores, 
                                    bins=30, 
                                    kde=True, 
                                    color='#2E86AB', 
                                    alpha=0.7,
                                    edgecolor='white',
                                    linewidth=1.2,
                                    kde_kws={'linewidth': 3, 'color': '#A23B72'}
                                )
                                
                                # 하위 20% 영역
                                ax_main.axvspan(0, percentile_20, alpha=0.25, color='#F18F01', label=f'하위 20% 영역')
                                ax_main.axvline(percentile_20, color='#F18F01', linestyle='-', linewidth=2.5, alpha=0.9)
                                ax_main.axvline(mean_score, color='#06A77D', linestyle='--', linewidth=2.5, alpha=0.9, label=f'평균: {mean_score:.1f}점')
                                ax_main.axvline(median_score, color='#D56062', linestyle='--', linewidth=2.5, alpha=0.9, label=f'중앙값: {median_score:.1f}점')
                                
                                ax_main.set_xlabel('총점', fontsize=12, fontweight='bold')
                                ax_main.set_ylabel('학생 수', fontsize=12, fontweight='bold')
                                ax_main.set_title('학생 점수 분포 통합 분석', fontsize=15, fontweight='bold', pad=15)
                                ax_main.legend(loc='upper right', fontsize=9)
                                ax_main.grid(True, alpha=0.3, linestyle='--')
                                ax_main.set_facecolor('#F8F9FA')
                                
                                # 박스플롯 (하단 왼쪽)
                                ax_box = fig.add_subplot(gs[2, 0])
                                bp = ax_box.boxplot(
                                    [scores],
                                    vert=True,
                                    patch_artist=True,
                                    boxprops=dict(facecolor='#2E86AB', alpha=0.7),
                                    medianprops=dict(color='#D56062', linewidth=2),
                                    whiskerprops=dict(color='#2C3E50', linewidth=1.5)
                                )
                                ax_box.axhline(percentile_20, color='#F18F01', linestyle='--', linewidth=1.5)
                                ax_box.set_ylabel('총점', fontsize=10, fontweight='bold')
                                ax_box.set_title('박스플롯', fontsize=11, fontweight='bold')
                                ax_box.grid(True, alpha=0.3, linestyle='--')
                                ax_box.set_facecolor('#F8F9FA')
                                
                                # 통계 요약 (하단 오른쪽)
                                ax_stats = fig.add_subplot(gs[2, 1])
                                ax_stats.axis('off')
                                
                                stats_text = f"""
                                📊 통계 요약
                                
                                전체 학생 수: {len(scores)}명
                                평균 점수: {mean_score:.2f}점
                                중앙값: {median_score:.2f}점
                                표준편차: {std_score:.2f}점
                                최고점: {scores.max():.1f}점
                                최저점: {scores.min():.1f}점
                                하위 20% 경계: {percentile_20:.1f}점
                                하위 20% 학생 수: {len(bottom_20_percent)}명
                                """
                                
                                ax_stats.text(
                                    0.1, 0.5, stats_text,
                                    fontsize=11,
                                    verticalalignment='center',
                                    fontfamily='monospace',
                                    bbox=dict(boxstyle='round', facecolor='#F8F9FA', alpha=0.8, edgecolor='#2C3E50', linewidth=1.5)
                                )
                                
                                fig.patch.set_facecolor('white')
                                st.pyplot(fig)
                                plt.close(fig)
                            
                            # 통계 정보 표시
                            col1, col2, col3, col4 = st.columns(4)
                            with col1:
                                st.metric("전체 학생 수", f"{len(scores)}명")
                            with col2:
                                st.metric("평균 점수", f"{mean_score:.1f}점")
                            with col3:
                                st.metric("중앙값", f"{median_score:.1f}점")
                            with col4:
                                st.metric("하위 20% 경계", f"{percentile_20:.1f}점")
                            
                            st.markdown("---")
                            
                            # 하위 20% 학생 리스트
                            st.subheader("⚠️ 하위 20% 학생 리스트")
                            
                            if len(bottom_20_percent) > 0:
                                # 학생명 열 찾기
                                student_column = None
                                for col in bottom_20_percent.columns:
                                    if '학생' in col:
                                        student_column = col
                                        break
                                
                                if student_column:
                                    # 하위 20% 학생 데이터 정리
                                    bottom_20_data = {
                                        "학생": bottom_20_percent[student_column].tolist(),
                                        "총점": bottom_20_percent[total_score_column].tolist()
                                    }
                                    
                                    # 총점 기준으로 정렬
                                    bottom_20_df = pd.DataFrame(bottom_20_data)
                                    bottom_20_df = bottom_20_df.sort_values('총점', ascending=True)
                                    bottom_20_df = bottom_20_df.reset_index(drop=True)
                                    
                                    # 순위 추가
                                    bottom_20_df.insert(0, '순위', range(1, len(bottom_20_df) + 1))
                                    
                                    # 표시
                                    st.info(f"💡 총 {len(bottom_20_df)}명의 학생이 하위 20%에 해당합니다. (총점 {percentile_20:.1f}점 이하)")
                                    
                                    # 데이터프레임 표시 (빨간색 강조)
                                    st.dataframe(
                                        bottom_20_df,
                                        use_container_width=True,
                                        hide_index=True
                                    )
                                    
                                    # 하위 20% 학생 이름만 리스트로 표시
                                    st.markdown("**하위 20% 학생 목록:**")
                                    student_list = bottom_20_df['학생'].tolist()
                                    student_list_text = ", ".join(student_list)
                                    st.markdown(f"*{student_list_text}*")
                                else:
                                    st.warning("⚠️ 엑셀 파일에서 '학생' 열을 찾을 수 없습니다.")
                            else:
                                st.info("하위 20%에 해당하는 학생이 없습니다.")
                            
                except Exception as e:
                    st.error(f"❌ 엑셀 파일 분석 중 오류 발생: {str(e)}")
                    st.info("💡 엑셀 파일 형식이 올바른지 확인해주세요. '점수 요약' 시트에 '학생'과 '총점' 열이 있어야 합니다.")

if __name__ == "__main__":
    main()

